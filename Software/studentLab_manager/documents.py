import docx

class ManagerDocuments():

    def load_pralabtemplate(doc):
        try:
            doc = docx.Document(doc)

            doc_readed = {}
            ping = 0
            previous = None
            add = []

            for para in doc.paragraphs:
                if para.style.name == 'Heading 1' and ping == 0:
                    previous = para.text.lower()
                    ping = 1
                elif para.style.name == 'Normal' or para.style.name == 'List Paragraph' and ping == 1:
                    add.append(para.text)
                elif para.style.name == 'Heading 1' and ping == 1:
                    doc_readed[previous] = add
                    add = []
                    previous = para.text.lower()
                    ping == 0
            doc_readed[previous] = add
            keys_indoc = doc_readed.keys()
            if not ("código" in keys_indoc and "nombre" in keys_indoc and 
                "área" in keys_indoc and "objetivos" in keys_indoc and 
                "materiales" in keys_indoc and "procedimientos" in keys_indoc and
                "preguntas" in keys_indoc):
                return 1
            elif not (len(doc_readed["código"]) > 0 and len(doc_readed["nombre"]) > 0 and
                len(doc_readed["área"]) > 0 and len(doc_readed["objetivos"]) > 0 and
                len(doc_readed["materiales"]) > 0 and len(doc_readed["preguntas"]) > 0) :
                return 2
            
            tables_size = len(doc.tables)
            if tables_size == 0:
                return 2
            elif tables_size > 1:
                return 1
            else:           
                for table in doc.tables:
                    if len(table.rows) > 1:
                        temp = [row for row in table.rows]
                        header = temp.pop(0).cells
                        headers = [hd.text for hd in header]
                        if len(headers) == 4:
                            if (headers[0] == "Descripción" and headers[1] == "Medidas" and
                                headers[2] == "Unidades" and headers[3] == "Tiempo"):
                                procs_data = []
                                for row in temp:
                                    cells = row.cells
                                    descp = cells[0].text
                                    raw_measu = cells[1].text.lower()
                                    measu = raw_measu.split(',')
                                    raw_units = cells[2].text.lower()
                                    units = raw_units.split(',')
                                    raw_time = cells[3].text.lower()
                                    time = raw_time.split(':')
                                    if(len(measu) == len(units)):
                                        print("xD")
                                    else:
                                        return 4

                                    if len(time) != 3:
                                        return 5

                                    try:
                                        new_proc = {'descripción':descp, 'medidas':measu, 'unidades':units,
                                            'tiempo': [int(t) for t in time]}
                                        procs_data.append(new_proc)
                                    except ValueError:
                                        return 5

                                try:
                                    doc_readed["código"] = int(doc_readed["código"][0])
                                    doc_readed["nombre"] = doc_readed["nombre"][0]
                                    doc_readed["área"] = doc_readed["área"][0]
                                    doc_readed["procedimientos"] = procs_data
                                    return doc_readed                            
                                except ValueError:
                                    return 6         
                            else:
                                return 3
                        else:
                            return 3
                    else:
                        return 1
        except IOError:
            print('There was an error opening the file!')
            return None