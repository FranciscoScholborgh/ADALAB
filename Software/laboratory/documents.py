from docx import Document
from utils.pydocx_utils import StyleDocument

class LabDocuments():

    def generate_preinforme(header, data, preguntas):
        practica = header["title"]
        curso = header["curso"]
        estudiantes = header["estudiantes"]
        id_grupo = header["id_grupo"]

        procedimientos = data["procedimientos"]
        proc_values = data["proc_data"]

        list_preguntas = preguntas["preguntas"]
        respuestas = preguntas["respuestas"]

        document = Document()

        style_header = StyleDocument.create_style(document, "style_header", is_bold=True, alignment='center')
        doc_header = document.add_heading(f'Práctica: {practica}', 0)
        doc_header.style = style_header
        
        style_subheader = StyleDocument.create_style(document, "style_subheader", alignment='center')
        curso_doc = document.add_heading(f"Curso: {curso.nombre}", 0)
        curso_doc.style = style_subheader
        docente_doc = document.add_heading(f"Docente: {curso.docente.nombre}", 0)
        docente_doc.style = style_subheader

        sec_header = document.add_heading(f'Integrantes del grupo #{id_grupo}', 0)
        sec_header.style = style_header

        for estudiante in estudiantes:
            estudiante_info= document.add_heading(f"{estudiante.nombre}, Código: {estudiante.codigo}", 0)
            estudiante_info.style = style_subheader

        style_title = StyleDocument.create_style(document, "style_title", is_bold=True)
        data_title = document.add_heading("Datos recolectados")
        data_title.style = style_title

        style_paragraph = StyleDocument.create_style(document, "style_paragraph", alignment='justify')

        data_index = 0
        for pasoPractica in procedimientos:
            paso = str(pasoPractica["paso"])
            try:
                proc_data = proc_values[paso]
                procedimiento = pasoPractica["procedimiento"]
                descrip = procedimiento.descripcion
                paso_style = document.add_paragraph('')
                paso_style.style = style_paragraph
                paso_style.add_run(f'Paso{paso}: ').italic = True
                paso_style.add_run(f'{descrip}')
                medidas = procedimiento.medida.split(",")
                unidades = procedimiento.unidad.split(",")
                max_col = len(medidas) + 1
                table = document.add_table(rows=1, cols=max_col)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Tiempo(s)'
                index = 1
                for medida in medidas:
                    hdr_cells[index].text = f'{medida}({unidades[index-1]})'
                    index+=1

                time = len(proc_data[0]) 
                for t in range(0, time):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(t)
                    len_medidas = len(proc_data)
                    for tab in range(0, len_medidas):
                        row_cells[tab+1].text = str(proc_data[tab][t])
            except KeyError:
                pass

        pregs_title = document.add_heading("Preguntas")
        pregs_title.style = style_title
        
        index = 0
        for pregunta in list_preguntas:
            preg = document.add_paragraph(f'{pregunta["text"]}', style='List Number')
            StyleDocument.editParagraph_style(preg)
            id_preg = str(pregunta["id"])
            respuesta = respuestas[id_preg]
            answ = document.add_paragraph(f'{respuesta}')
            answ.style = style_paragraph
            index += 1

        return document